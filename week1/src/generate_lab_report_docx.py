from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
REPORT_DIR = ROOT.parent / "report"
REPORT_MD = REPORT_DIR / "实验1-空间域点运算_实验报告.md"
REPORT_DOCX = REPORT_DIR / "实验1-空间域点运算_实验报告.docx"


def set_run_font(run, size: int = 11, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_text(cell, text: str, bold: bool = False, align_center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_code_block(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_text(cell, code.rstrip())
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    cell._tc.get_or_add_tcPr().append(shading)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)


_INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_markdown_text(paragraph, text: str) -> None:
    pos = 0
    for match in _INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_image(document: Document, alt: str, image_path: Path) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, after=3)
    run = caption.add_run(alt)
    set_run_font(run, size=10, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=8)
    run = paragraph.add_run()
    width = Inches(6.2)
    if image_path.name == "08_all_results.png":
        width = Inches(6.4)
    elif image_path.name == "07_histogram_comparison.png":
        width = Inches(6.3)
    run.add_picture(str(image_path), width=width)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, cell_text, bold=(r_idx == 0), align_center=True)
            if r_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E8EEF7")
                cell._tc.get_or_add_tcPr().append(shading)
    document.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    inner = line.strip().strip("|")
    return [cell.strip() for cell in inner.split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def add_centered_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=6, after=14)
    run = paragraph.add_run(text)
    set_run_font(run, size=22, bold=True)


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    title_emitted = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines))
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Markdown table: pipe-prefixed line followed by separator row
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(stripped)]
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(parse_table_row(lines[j]))
                j += 1
            add_markdown_table(document, rows)
            i = j
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            alt, rel_path = image_match.groups()
            add_image(document, alt, (REPORT_MD.parent / rel_path).resolve())
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            heading_text = heading_match.group(2)
            if level == 1 and not title_emitted:
                add_centered_title(document, heading_text)
                title_emitted = True
            else:
                paragraph = document.add_heading(level=level)
                run = paragraph.add_run(heading_text)
                set_run_font(run, size=15 if level == 2 else 13, bold=True)
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_markdown_text(paragraph, stripped[2:])
            set_paragraph_spacing(paragraph, after=2)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_markdown_text(paragraph, numbered.group(1))
            set_paragraph_spacing(paragraph, after=2)
            i += 1
            continue

        paragraph = document.add_paragraph()
        add_markdown_text(paragraph, stripped)
        set_paragraph_spacing(paragraph)
        i += 1

    document.save(REPORT_DOCX)
    print(REPORT_DOCX)


if __name__ == "__main__":
    build_docx()
